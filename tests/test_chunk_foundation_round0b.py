"""Round 0B regression tests for filtering and DOCX heading fallback."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt
from langchain_core.documents import Document

from rag_knowledge.services.chunk_health_audit import (
    diagnose_low_information,
    is_suspicious_heading,
)
from rag_knowledge.services.loader import FileLoader
from rag_knowledge.services.unstructured_loader import UnstructuredChapterLoader


ROUND0A_DIR = Path("docs/04_已完成归档/01_文档解析与切块/已完成-第0A轮-基线审计与问题定位/审计产物")


def _paragraph(text: str, *, size: int = 12, bold: bool = False):
    document = DocxDocument()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return paragraph


def test_low_information_keeps_all_round0a_false_positive_regressions():
    samples = json.loads(
        (ROUND0A_DIR / "filter_false_positive_regression.json").read_text(
            encoding="utf-8"
        )
    )

    rejected = [
        sample["id"]
        for sample in samples
        if FileLoader._is_low_information(sample["text"])
    ]

    assert rejected == []


def test_low_information_still_rejects_long_garbled_run():
    garbled = (
        "Description=TongWeb Server\n"
        "After=database.target\n"
        "[Servi散൝祔数昽牯楫杮唍敳㵲潲瑯䔍癮物湯敭瑮∽䅊䅖䡟䵏㵅"
        "䅊䅖䡟䵏彅䅖䥒呁≅倍䑉楆敬\n"
        "执行/data/tong/TongWeb7/bin目录下的installservice.sh\n"
    )

    assert FileLoader._is_low_information(garbled) is True


def test_audit_diagnosis_marks_unexpected_script_as_garbled():
    garbled = "配置服务并执行启动脚本൝，随后运行 installservice.sh。"

    decision = diagnose_low_information(garbled)

    assert decision.rejected is True
    assert decision.reason_code == "unexpected_script"


def test_low_information_keeps_normal_technical_math_symbol():
    text = "当阈值≥0.95 时，使用 WebGL2 与 Redis 服务执行增量发布。"

    assert FileLoader._is_low_information(text) is False


def test_audit_diagnosis_keeps_chinese_technical_mixed_content():
    text = (
        "配置 Turnserver：设置 listening-port、tls-listening-port、fingerprint、realm，"
        "并写入 user 凭证后重启服务。"
    )

    decision = diagnose_low_information(text)

    assert decision.rejected is False
    assert decision.reason_code == "keep"


def test_12pt_body_is_not_promoted_to_heading_without_style_or_numbering():
    paragraph = _paragraph("配置 Redis 与 Cesium/WebGL 后重启服务。", size=12)

    assert UnstructuredChapterLoader._is_conservative_heading_fallback(paragraph) is None


def test_command_config_and_port_are_not_fallback_headings():
    for text in (
        "vim /etc/sysctl.conf",
        "enabled=1",
        "5349：TLS/TCP，TLS服务",
    ):
        paragraph = _paragraph(text, size=16, bold=True)
        assert UnstructuredChapterLoader._is_conservative_heading_fallback(paragraph) is None


def test_numbered_heading_remains_fallback_heading():
    paragraph = _paragraph("1.2 安装与部署", size=12)

    assert UnstructuredChapterLoader._is_conservative_heading_fallback(paragraph) == 2


def test_docx_cover_fallback_does_not_create_section_path(tmp_path):
    document = DocxDocument()
    cover = document.add_paragraph()
    run = cover.add_run("2024年5月")
    run.font.size = Pt(16)
    run.bold = True
    document.add_paragraph("这是封面后的说明文字。")
    path = tmp_path / "manual.docx"
    document.save(path)

    docs = UnstructuredChapterLoader().load(str(path))

    assert docs[0].metadata["section_path"] == ""


def test_markdown_toc_is_not_a_section_or_retrievable_body(tmp_path):
    class _Element:
        def __init__(self, category: str, text: str):
            self.category = category
            self.text = text

        def __str__(self) -> str:
            return self.text

    elements = [
        _Element("Title", "文章目录"),
        _Element("ListItem", "1. MySQL 数据类型"),
        _Element("Title", "1.MySQL 数据类型"),
        _Element("NarrativeText", "正文内容。"),
    ]
    path = tmp_path / "article.md"
    path.write_text("fixture", encoding="utf-8")

    docs = UnstructuredChapterLoader()._parse_unstructured(path, lambda **_: elements)

    assert len(docs) == 1
    assert docs[0].section_path == ["1.MySQL 数据类型"]
    assert "目录" not in docs[0].content_markdown


def test_cd_dvd_setting_is_not_a_command_like_audit_heading():
    suspicious, reason = is_suspicious_heading("CD/DVD设置")

    assert suspicious is False
    assert reason == ""


def test_unstructured_loader_keeps_consecutive_title_only_items(tmp_path):
    class _Element:
        def __init__(self, category: str, text: str):
            self.category = category
            self.text = text

        def __str__(self) -> str:
            return self.text

    elements = [
        _Element("UncategorizedText", "1.陕西地信================"),
        _Element("Title", "2.资源目录图层树名称抽取"),
        _Element("Title", "3.专题数据图层树名称抽取"),
        _Element("Title", "专业控制台"),
        _Element("Title", "1.图层控制"),
        _Element("NarrativeText", "支持瓦片配色和图例。"),
    ]
    path = tmp_path / "features.txt"
    path.write_text("fixture", encoding="utf-8")

    docs = UnstructuredChapterLoader()._parse_unstructured(path, lambda **_: elements)

    body = next(element for element in docs if element.content_type == "text")
    assert body.section_path == ["1.陕西地信"]
    assert "2.资源目录图层树名称抽取" in body.content_markdown
    assert "3.专题数据图层树名称抽取" in body.content_markdown
    assert "专业控制台" in body.content_markdown
    assert "支持瓦片配色" in body.content_markdown


def test_post_process_drops_section_prefixed_toc_marker():
    loader = object.__new__(FileLoader)
    chunks = [
        Document(
            page_content="# 2024年5月\n\n目录",
            metadata={"content_type": "text"},
        )
    ]

    assert loader._post_process_chunks(chunks) == []
