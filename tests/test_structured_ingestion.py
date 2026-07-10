import importlib.util
import json
import sys
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest
from docx import Document as WordDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from rag_knowledge.services.bm25_store import BM25Store
from rag_knowledge.services.retrieval_quality import RetrievalQualityStrategy
from rag_knowledge.repository.relational_db import RelationalDB
from tests.fixtures.pipeline_graph_facts import seed_pipeline_table_graph


PROJECT_ROOT = Path(__file__).resolve().parents[1]


@pytest.fixture(autouse=True)
def _graph_storage(isolated_storage):
    import shutil
    _, _, _, data_dir = isolated_storage(
        db_name="structured-ingestion.db",
        data_dir_name="structured-ingestion-data",
    )
    policies_src = PROJECT_ROOT / "data" / "retrieval_intent_policies.json"
    shutil.copy2(policies_src, data_dir / "retrieval_intent_policies.json")
    seed_pipeline_table_graph(RelationalDB())


def _load_module(module_name: str, relative_path: str):
    path = PROJECT_ROOT / relative_path
    spec = importlib.util.spec_from_file_location(module_name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module


def _load_real_loader_module():
    real_unstructured = _load_module(
        "tests_real_unstructured_loader",
        "rag_knowledge/services/unstructured_loader.py",
    )
    loader_path = PROJECT_ROOT / "rag_knowledge/services/loader.py"
    spec = importlib.util.spec_from_file_location("tests_real_loader", loader_path)
    loader_module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    with patch.dict(sys.modules, {"rag_knowledge.services.unstructured_loader": real_unstructured}):
        spec.loader.exec_module(loader_module)
    return loader_module


def _create_file_loader(chunk_size: int = 500):
    loader_module = _load_real_loader_module()
    FileLoader = loader_module.FileLoader
    real_unstructured = _load_module(
        "tests_real_unstructured_loader_for_instance",
        "rag_knowledge/services/unstructured_loader.py",
    )
    loader = object.__new__(FileLoader)
    loader._chunk_size = chunk_size
    loader._chunk_overlap = 0
    loader._semantic_chunking_enabled = False
    loader._semantic_chunker = None
    loader._extract_images = False
    loader._use_unstructured = True
    loader._splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=0,
        separators=["\n\n", "\n", "。", "；", "，", " ", ""],
    )
    loader._unstructured_loader = real_unstructured.UnstructuredChapterLoader(
        chunk_size=chunk_size,
        chunk_overlap=0,
        strategy="fast",
    )
    return loader


class StructuredDocxIngestionTests(unittest.TestCase):
    def _build_docx(self, builder) -> Path:
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)
        path = Path(temp_dir.name) / "structured.docx"
        doc = WordDocument()
        builder(doc)
        doc.save(path)
        return path

    def test_docx_table_chunk_uses_markdown_and_full_section_path(self):
        path = self._build_docx(self._build_pipeline_table_doc)
        loader = _create_file_loader()

        chunks = loader._load_text(str(path))

        table_chunks = [chunk for chunk in chunks if chunk.metadata.get("content_type") == "table"]
        self.assertEqual(len(table_chunks), 1)
        table_chunk = table_chunks[0]
        self.assertTrue(
            table_chunk.page_content.startswith(
                "# PipelineBuilder > 数据规范 > 管线点表 > 点数据结构"
            )
        )
        self.assertIn("| 字段名 | 说明 |", table_chunk.page_content)
        self.assertIn("| 管点编号 | 唯一识别码，全局唯一，必要字段 |", table_chunk.page_content)
        self.assertEqual(
            table_chunk.metadata.get("section_path"),
            "PipelineBuilder > 数据规范 > 管线点表 > 点数据结构",
        )
        self.assertEqual(table_chunk.metadata.get("chunking_method"), "table")
        self.assertIn("PipelineBuilder", table_chunk.metadata.get("searchable_text", ""))
        self.assertIn("管点编号", table_chunk.metadata.get("searchable_text", ""))

    def test_docx_heading_boundary_forces_section_split(self):
        def build(doc):
            doc.add_heading("服务 A", level=2)
            doc.add_paragraph("这是 A 的说明。")
            doc.add_heading("服务 B", level=2)
            doc.add_paragraph("这是 B 的说明。")

        path = self._build_docx(build)
        loader = _create_file_loader()

        chunks = loader._load_text(str(path))

        text_chunks = [chunk for chunk in chunks if chunk.metadata.get("content_type") != "table"]
        section_to_content = {
            chunk.metadata.get("section_path"): chunk.page_content for chunk in text_chunks
        }
        self.assertIn("服务 A", section_to_content)
        self.assertIn("服务 B", section_to_content)
        self.assertIn("这是 A 的说明", section_to_content["服务 A"])
        self.assertNotIn("服务 B", section_to_content["服务 A"])
        self.assertIn("这是 B 的说明", section_to_content["服务 B"])
        self.assertNotIn("服务 A", section_to_content["服务 B"])

    @staticmethod
    def _build_pipeline_table_doc(doc: WordDocument) -> None:
        doc.add_heading("PipelineBuilder", level=1)
        doc.add_heading("数据规范", level=2)
        doc.add_heading("管线点表", level=3)
        doc.add_heading("点数据结构", level=4)
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "字段名"
        table.cell(0, 1).text = "说明"
        table.cell(1, 0).text = "管点编号"
        table.cell(1, 1).text = "唯一识别码，全局唯一，必要字段"
        table.cell(2, 0).text = "地面高程"
        table.cell(2, 1).text = "管线点地表面的高程，单位为米，必要字段"


class RetrievalQualityTableBoostTests(unittest.TestCase):
    def test_table_queries_boost_table_chunks(self):
        cfg = SimpleNamespace(
            retrieval_quality=SimpleNamespace(
                enabled=True,
                score_threshold_enabled=False,
                score_threshold=0.35,
                jaccard_dedup_enabled=False,
                jaccard_threshold=0.85,
                dynamic_topk_enabled=False,
                score_drop_ratio=0.5,
                min_top_k=3,
                max_top_k=8,
                contextual_compression_enabled=False,
                debug_log_enabled=False,
            )
        )
        strategy = RetrievalQualityStrategy(cfg)
        docs = [
            Document(
                page_content="管线发布服务配置需要修改 Apache 和服务目录。",
                metadata={"chunk_id": "service", "rrf_score": 0.45},
            ),
            Document(
                page_content=(
                    "# PipelineBuilder > 数据规范 > 管线点表 > 点数据结构\n\n"
                    "| 字段名 | 说明 |\n| --- | --- |\n| 管点编号 | 唯一识别码 |"
                ),
                metadata={
                    "chunk_id": "table",
                    "rrf_score": 0.40,
                    "content_type": "table",
                    "chunking_method": "table",
                },
            ),
        ]

        result = strategy.apply("管线点表规范", docs)

        self.assertEqual(result[0].metadata.get("chunk_id"), "table")
        self.assertGreater(
            float(result[0].metadata.get("quality_score", 0.0)),
            float(result[1].metadata.get("quality_score", 0.0)),
        )


class Bm25SearchableTextTests(unittest.TestCase):
    def test_bm25_prefers_searchable_text_metadata_when_present(self):
        store = object.__new__(BM25Store)
        store._bm25 = None
        store._docs = []
        store._metadatas = []
        chroma = MagicMock()
        chroma.get.return_value = {
            "documents": [
                "| 字段名 | 说明 |\n| --- | --- |\n| 管点编号 | 唯一识别码 |"
            ],
            "metadatas": [
                {
                    "source": "StampTools用户手册.docx",
                    "searchable_text": (
                        "PipelineBuilder 数据规范 管线 点表 规范 点数据结构 "
                        "字段名 说明 管点编号 唯一识别码"
                    ),
                }
            ],
            "ids": ["chunk-1"],
        }
        vector_store = MagicMock()
        vector_store.get_chroma.return_value = chroma

        with patch("rag_knowledge.services.bm25_store.VectorStore", return_value=vector_store):
            store._build_index()

        result = store.search("管线点表规范", review_status=None, top_k=1)

        self.assertEqual([doc.metadata.get("chunk_id") for doc in result], ["chunk-1"])


class StructuredRegressionDatasetTests(unittest.TestCase):
    def test_structured_retrieval_regression_dataset_exists_with_expected_cases(self):
        dataset_path = PROJECT_ROOT / "data/structured_retrieval_regression.json"
        self.assertTrue(dataset_path.exists())

        data = json.loads(dataset_path.read_text(encoding="utf-8"))
        questions = {item["question"] for item in data}

        self.assertIn("管线点表规范", questions)
        self.assertIn("发布管线点表的要求", questions)
        self.assertIn("PipelineBuilder 管线点表字段要求", questions)
        self.assertIn("PipelineBuilder 如何发布管线", questions)
        self.assertIn("管线发布服务如何配置", questions)
        self.assertIn("DOMBuilder 工具是影像工具吗", questions)
        self.assertIn("DOMBuilder 如何发布影像", questions)

        target = next(item for item in data if item["question"] == "管线点表规范")
        self.assertEqual(target["expected"][0]["source"], "StampTools用户手册.docx")
        self.assertIn("PipelineBuilder > 数据规范 > 管线点表", target["expected"][0]["section_path_contains"])
        self.assertIn("管点编号", target["expected"][0]["content_must_include"])


if __name__ == "__main__":
    unittest.main()
