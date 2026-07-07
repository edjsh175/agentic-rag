import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace

from docx import Document as WordDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from rag_knowledge.services.loader import FileLoader
from rag_knowledge.services.retrieval_quality import RetrievalQualityStrategy
from rag_knowledge.services.unstructured_loader import UnstructuredChapterLoader


def _build_loader(chunk_size: int = 220) -> FileLoader:
    loader = object.__new__(FileLoader)
    loader._chunk_size = chunk_size
    loader._chunk_overlap = 0
    loader._semantic_chunking_enabled = False
    loader._semantic_chunker = None
    loader._extract_images = False
    loader._use_unstructured = True
    loader._splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=0,
        separators=["\n\n", "\n", "。", "；", "，", " ", ""],
    )
    loader._unstructured_loader = UnstructuredChapterLoader(
        chunk_size=chunk_size,
        chunk_overlap=0,
        strategy="fast",
    )
    return loader


class DocxTableGuardsTests(unittest.TestCase):
    def _build_docx(self) -> Path:
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)
        path = Path(temp_dir.name) / "pipeline.docx"
        doc = WordDocument()
        doc.add_heading("PipelineBuilder", level=1)
        doc.add_heading("数据规范", level=2)
        doc.add_heading("管线点表", level=3)

        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "字段名"
        table.cell(0, 1).text = "说明"
        rows = [
            ("管点编号", "唯一识别码，全局唯一，必要字段"),
            ("地面高程", "管线点地表面的高程，单位为米，必要字段"),
            ("特征", "多通点、转折点、偏心点等，必要字段"),
            ("附属物设施", "检修井、阀门井等，必要字段"),
            ("偏心井点号", "偏心点时填入井盖中心点的管点编号"),
            ("离心井点号", "离心井时填入井盖中心点的管点编号"),
            ("旋转角度", "正东方向或正北方向的旋转角度"),
            ("井类型", "圆井、方井、台井、离心井等"),
        ]
        for field_name, description in rows:
            cells = table.add_row().cells
            cells[0].text = field_name
            cells[1].text = description

        doc.save(path)
        return path

    def test_docx_table_keeps_single_chunk_under_soft_limit(self):
        loader = _build_loader(chunk_size=220)
        path = self._build_docx()

        chunks = loader._load_text(str(path))
        table_chunks = [
            chunk for chunk in chunks
            if chunk.metadata.get("section_path") == "PipelineBuilder > 数据规范 > 管线点表"
            and chunk.metadata.get("content_type") == "table"
        ]

        self.assertEqual(len(table_chunks), 1)
        table_chunk = table_chunks[0]
        self.assertIn("| 管点编号 | 唯一识别码，全局唯一，必要字段 |", table_chunk.page_content)
        self.assertIn("| 井类型 | 圆井、方井、台井、离心井等 |", table_chunk.page_content)
        self.assertEqual(table_chunk.metadata.get("row_start"), 1)
        self.assertEqual(table_chunk.metadata.get("row_end"), 8)


class TableQueryBoostGuardsTests(unittest.TestCase):
    def test_real_query_hints_apply_table_bonus(self):
        cfg = SimpleNamespace(
            retrieval_quality=SimpleNamespace(
                enabled=True,
                score_threshold_enabled=False,
                score_threshold=0.35,
                jaccard_dedup_enabled=False,
                jaccard_threshold=0.85,
                dynamic_topk_enabled=False,
                score_drop_ratio=0.5,
                min_top_k=3,
                max_top_k=8,
                contextual_compression_enabled=False,
                debug_log_enabled=False,
            )
        )
        strategy = RetrievalQualityStrategy(cfg)
        docs = [
            Document(
                page_content="管线发布服务配置需要修改 Apache 和服务目录。",
                metadata={"chunk_id": "service", "rrf_score": 0.45},
            ),
            Document(
                page_content=(
                    "# PipelineBuilder > 数据规范 > 管线点表\n\n"
                    "| 字段名 | 说明 |\n"
                    "| --- | --- |\n"
                    "| 管点编号 | 唯一识别码 |\n"
                    "| 地面高程 | 管线点地表面的高程 |\n"
                ),
                metadata={
                    "chunk_id": "table",
                    "rrf_score": 0.40,
                    "content_type": "table",
                    "chunking_method": "table",
                },
            ),
        ]

        result = strategy.apply("管线点表规范", docs)

        self.assertEqual(result[0].metadata.get("chunk_id"), "table")
        self.assertGreater(result[0].metadata.get("table_query_boost", 0.0), 0.0)


if __name__ == "__main__":
    unittest.main()
