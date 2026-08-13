"""Acceptance probe: can WebRTC/WebGL DOCX manuals support map style L1 coding?

Two layers:
1) Offline DOCX coverage (heading + keyword evidence for line/polygon/style APIs)
2) Live Hybrid retrieval hit-rate on curated acceptance questions (optional if chroma open)

Usage (project root, venv):
  .\\venv\\Scripts\\python.exe scripts\\validate_map_style_l1_feasibility.py
  .\\venv\\Scripts\\python.exe scripts\\validate_map_style_l1_feasibility.py --skip-retrieval
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

DOCX_FILES = [
    ROOT / "watch_directory" / "word" / "StampGIS平台WebRTC接口说明书_2026_05_26_15_45_50.docx",
    ROOT / "watch_directory" / "word" / "StampGIS平台WebGL接口说明书_2026_05_26_15_45_46.docx",
]

SOURCE_NAME_HINTS = (
    "WebRTC接口说明书",
    "WebGL接口说明书",
    "GIS平台WebRTC接口说明书",  # legacy filename variant in chroma
)

# Acceptance questions for map-related styling in secondary development.
ACCEPTANCE_CASES = [
    {
        "id": "style-line-color",
        "question": "StampGIS 二次开发里怎么设置地图折线或线的颜色？",
        "must_any": ["线", "颜色", "折线", "polyline", "Polyline", "覆盖物"],
        "nice_any": ["setColor", "线宽", "样式", "style"],
    },
    {
        "id": "style-polygon-fill",
        "question": "StampGIS 二次开发里怎么设置面或多边形的填充颜色？",
        "must_any": ["面", "多边形", "填充", "polygon", "Polygon", "覆盖物", "颜色"],
        "nice_any": ["透明度", "边线", "轮廓"],
    },
    {
        "id": "style-overlay-color",
        "question": "怎么设置覆盖物颜色？",
        "must_any": ["覆盖物", "颜色"],
        "nice_any": ["示例", "StampUtil", "function", "("],
    },
    {
        "id": "style-transparency",
        "question": "怎么设置地图覆盖物或模型的透明度？",
        "must_any": ["透明"],
        "nice_any": ["覆盖物", "模型", "地形", "瓦片"],
    },
    {
        "id": "style-line-width",
        "question": "二次开发如何修改线宽？",
        "must_any": ["线宽", "宽度", "线"],
        "nice_any": ["样式", "style", "折线"],
    },
    {
        "id": "init-vue-stamputil",
        "question": "Vue3 项目中如何引入 StampUtil 并初始化地图？",
        "must_any": ["StampUtil", "Vue", "import"],
        "nice_any": ["div", "初始化", "底图"],
    },
    {
        "id": "draw-polyline",
        "question": "如何在地图上绘制折线？",
        "must_any": ["折线", "画线", "绘制", "polyline", "Polyline", "线"],
        "nice_any": ["坐标", "示例", "add"],
    },
    {
        "id": "draw-polygon",
        "question": "如何在地图上绘制面或多边形？",
        "must_any": ["多边形", "画面", "绘制", "polygon", "Polygon", "面"],
        "nice_any": ["坐标", "示例", "add"],
    },
]


HEADING_STYLE_HINTS = re.compile(
    r"颜色|样式|覆盖物|线宽|折线|多边形|填充|透明|边线|轮廓|面|线|"
    r"polyline|polygon|symbol|style|color|width|fill",
    re.I,
)
BODY_STYLE_HINTS = re.compile(
    r"设置覆盖物颜色|线宽|线颜色|面颜色|填充色|边线|轮廓线|"
    r"折线|多边形|polyline|polygon|setColor|setStyle|透明度|覆盖物颜色",
    re.I,
)


@dataclass
class DocxCoverage:
    file: str
    exists: bool
    size: int = 0
    nonempty_paras: int = 0
    headings_total: int = 0
    style_related_headings: list[str] = field(default_factory=list)
    style_related_body_samples: list[str] = field(default_factory=list)
    keyword_counts: dict[str, int] = field(default_factory=dict)
    codeish_paras: int = 0
    vue_stamputil_hits: int = 0


def _scan_docx(path: Path) -> DocxCoverage:
    from docx import Document

    cov = DocxCoverage(file=path.name, exists=path.exists(), size=path.stat().st_size if path.exists() else 0)
    if not path.exists():
        return cov
    doc = Document(str(path))
    keys = [
        "颜色", "样式", "覆盖物", "线宽", "折线", "多边形", "填充", "透明",
        "边线", "轮廓", "StampUtil", "Vue", "示例", "代码", "polyline", "polygon",
        "setColor", "setStyle", "style", "width", "fill",
    ]
    full_parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        cov.nonempty_paras += 1
        full_parts.append(text)
        style = getattr(getattr(para, "style", None), "name", "") or ""
        if "Heading" in style or "标题" in style:
            cov.headings_total += 1
            if HEADING_STYLE_HINTS.search(text):
                cov.style_related_headings.append(f"{style}: {text[:160]}")
        elif BODY_STYLE_HINTS.search(text):
            if len(cov.style_related_body_samples) < 40:
                cov.style_related_body_samples.append(text[:200])
        if any(tok in text for tok in ("function", "import ", "const ", "var ", "=>", "StampUtil")) and (
            "(" in text or "{" in text or "import " in text
        ):
            cov.codeish_paras += 1
        if "StampUtil" in text or "Vue" in text:
            cov.vue_stamputil_hits += 1
    full = "\n".join(full_parts)
    cov.keyword_counts = {k: full.count(k) for k in keys if full.count(k)}
    return cov


def _text_has_any(text: str, needles: list[str]) -> bool:
    low = text.lower()
    for n in needles:
        if n.lower() in low:
            return True
    return False


def _score_hit(text: str, case: dict) -> dict:
    must = _text_has_any(text, case["must_any"])
    nice = _text_has_any(text, case.get("nice_any") or [])
    has_code = bool(re.search(r"(import\s+|function\s+|=>|\(.*\)\s*;|StampUtil\.)", text))
    return {
        "must_hit": must,
        "nice_hit": nice,
        "code_snippet_like": has_code,
        "pass": must,  # minimum bar: must keywords present in evidence
    }


def _best_windows(text: str, needles: list[str], *, window: int = 700, limit: int = 5) -> list[str]:
    """Pick evidence windows preferring API/code body over TOC-like lines."""
    low = text.lower()
    cands: list[tuple[int, str]] = []
    for kw in needles:
        start = 0
        key = kw.lower()
        while True:
            idx = low.find(key, start)
            if idx < 0:
                break
            a = max(0, idx - 180)
            b = min(len(text), idx + window)
            snippet = text[a:b]
            score = 0
            if "StampUtil" in snippet:
                score += 5
            if re.search(r"(function|import\s+|参数|代码示例|fillColor|lineColor|lineWidth|width)", snippet, re.I):
                score += 3
            # penalize pure TOC / page-number rows
            if re.search(r"\t\d+\s*$", snippet.splitlines()[0] if snippet.splitlines() else ""):
                score -= 3
            if snippet.count("\t") >= 2 and "StampUtil" not in snippet:
                score -= 2
            cands.append((score, snippet))
            start = idx + len(key)
            if len(cands) > 80:
                break
    cands.sort(key=lambda x: x[0], reverse=True)
    out: list[str] = []
    seen = set()
    for score, snippet in cands:
        sig = snippet[:80]
        if sig in seen:
            continue
        seen.add(sig)
        out.append(snippet)
        if len(out) >= limit:
            break
    return out


def run_offline_case_match(coverages: list[DocxCoverage]) -> list[dict]:
    """Match acceptance questions against concatenated DOCX text (upper-bound feasibility)."""
    from docx import Document

    blobs: list[tuple[str, str]] = []
    for path, cov in zip(DOCX_FILES, coverages):
        if not cov.exists:
            continue
        doc = Document(str(path))
        text = "\n".join((p.text or "") for p in doc.paragraphs)
        blobs.append((path.name, text))

    results = []
    for case in ACCEPTANCE_CASES:
        best = None
        for name, text in blobs:
            hits = _best_windows(text, case["must_any"] + list(case.get("nice_any") or []))
            window = "\n---\n".join(hits) if hits else ""
            scored = _score_hit(window or text[:2000], case)
            scored["source"] = name
            scored["evidence_preview"] = (window or "")[:500]
            if best is None or (scored["pass"] and not best["pass"]) or (
                scored["pass"] == best["pass"]
                and (
                    (scored["code_snippet_like"] and not best["code_snippet_like"])
                    or (scored["nice_hit"] and not best["nice_hit"])
                )
            ):
                best = scored
        results.append({"id": case["id"], "question": case["question"], "offline": best or {"pass": False}})
    return results


def run_live_retrieval(top_k: int = 5) -> list[dict]:
    from rag_knowledge.services.retrieval_strategy import RetrievalStrategy

    strategy = RetrievalStrategy()
    out = []
    for case in ACCEPTANCE_CASES:
        docs = strategy.retrieve(
            case["question"],
            kb_name=None,
            review_status="approved",
            method="hybrid",
            top_k=top_k,
        )
        ranked = []
        for i, doc in enumerate(docs, 1):
            meta = doc.metadata or {}
            source = str(meta.get("source") or meta.get("file_name") or "")
            body = doc.page_content or ""
            scored = _score_hit(body, case)
            in_target_manual = any(h in source for h in SOURCE_NAME_HINTS)
            ranked.append(
                {
                    "rank": i,
                    "source": source,
                    "section_path": meta.get("section_path"),
                    "content_type": meta.get("content_type"),
                    "in_target_manual": in_target_manual,
                    "preview": body[:350],
                    **scored,
                }
            )
        target_hits = [r for r in ranked if r["in_target_manual"] and r["must_hit"]]
        any_must = [r for r in ranked if r["must_hit"]]
        out.append(
            {
                "id": case["id"],
                "question": case["question"],
                "retrieval": {
                    "returned": len(ranked),
                    "pass_target_manual": bool(target_hits),
                    "pass_any_approved": bool(any_must),
                    "best_target_rank": target_hits[0]["rank"] if target_hits else None,
                    "hits": ranked,
                },
            }
        )
    return out


def judge(offline_results: list[dict], retrieval_results: list[dict] | None, coverages: list[DocxCoverage]) -> dict:
    offline_pass = sum(1 for r in offline_results if (r.get("offline") or {}).get("pass"))
    offline_total = len(offline_results)
    retr_pass = None
    retr_total = None
    if retrieval_results is not None:
        retr_pass = sum(1 for r in retrieval_results if (r.get("retrieval") or {}).get("pass_target_manual"))
        retr_total = len(retrieval_results)

    style_heading_n = sum(len(c.style_related_headings) for c in coverages)
    style_api_evidence = any(
        any(
            key in h
            for key in (
                "创建折线",
                "创建多边形",
                "批量修改样式",
                "获取样式",
                "设置矢量楼块图层样式",
                "覆盖物颜色",
                "设置高亮颜色",
            )
            for h in c.style_related_headings
        )
        or any(key in s for key in ("fillColor", "lineColor", "lineWidth", "覆盖物颜色") for s in c.style_related_body_samples)
        for c in coverages
    )
    offline_code_pass = sum(1 for r in offline_results if (r.get("offline") or {}).get("code_snippet_like"))

    # Feasibility levels
    # A: manuals contain style APIs + offline cases mostly pass
    # B: content exists but retrieval to target manuals weak
    # C: content gap for line/polygon styling
    if offline_pass / max(offline_total, 1) >= 0.75 and style_api_evidence:
        level = "FEASIBLE_WITH_MANUALS"
        summary = (
            "两份接口手册对地图线/面/样式类二次开发具备内容基础（含创建折线/多边形、样式与颜色参数）；"
            "L1 可不依赖 Vue 整仓源码；完整可粘贴代码质量仍受 DOCX 切块影响。"
        )
    elif offline_pass / max(offline_total, 1) >= 0.5:
        level = "PARTIALLY_FEASIBLE"
        summary = (
            "手册覆盖部分地图样式能力，完整线/面样式问答仍有缺口；"
            "建议补少量 Cookbook 样例，而非整仓 Vue 源码。"
        )
    else:
        level = "NOT_FEASIBLE_WITH_MANUALS_ONLY"
        summary = "仅凭这两份手册难以稳定支撑地图线/面样式二次开发代码生成。"

    if retrieval_results is not None and retr_total:
        ratio = retr_pass / retr_total
        if level == "FEASIBLE_WITH_MANUALS" and ratio < 0.5:
            level = "CONTENT_OK_RETRIEVAL_WEAK"
            summary += f" 检索落点偏弱（目标手册命中 {retr_pass}/{retr_total}），入库切块/分类/审核状态可能影响可用性。"
        elif ratio >= 0.5:
            summary += f" 检索侧目标手册命中 {retr_pass}/{retr_total}。"

    return {
        "level": level,
        "summary": summary,
        "offline_pass": offline_pass,
        "offline_total": offline_total,
        "offline_code_like_pass": offline_code_pass,
        "retrieval_pass_target": retr_pass,
        "retrieval_total": retr_total,
        "style_related_heading_count": style_heading_n,
        "has_style_api_evidence": style_api_evidence,
        "vue_source_required_for_map_style": False
        if level.startswith("FEASIBLE") or level.startswith("CONTENT_OK")
        else "uncertain",
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--skip-retrieval", action="store_true")
    parser.add_argument("--top-k", type=int, default=5)
    parser.add_argument(
        "--output",
        type=Path,
        default=ROOT / "data" / "eval_map_style_l1_feasibility.json",
    )
    args = parser.parse_args()

    coverages = [_scan_docx(p) for p in DOCX_FILES]
    offline = run_offline_case_match(coverages)
    retrieval = None
    retrieval_error = None
    if not args.skip_retrieval:
        try:
            retrieval = run_live_retrieval(top_k=args.top_k)
        except Exception as exc:  # noqa: BLE001 - acceptance probe must continue
            retrieval_error = f"{type(exc).__name__}: {exc}"

    # merge offline + retrieval by id
    merged = []
    retr_by_id = {r["id"]: r for r in (retrieval or [])}
    for row in offline:
        item = dict(row)
        if row["id"] in retr_by_id:
            item["retrieval"] = retr_by_id[row["id"]]["retrieval"]
        merged.append(item)

    verdict = judge(offline, retrieval, coverages)
    report = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "docs": [asdict(c) for c in coverages],
        "verdict": verdict,
        "cases": merged,
        "retrieval_error": retrieval_error,
        "notes": [
            "offline=文档全文关键词窗口匹配，表示内容上界（手册里有没有这类信息）",
            "retrieval=Hybrid 检索 top_k 是否命中目标手册且含 must 关键词，表示系统可用性",
            "地图线/面样式属于 SDK API，不等于 Vue UI 皮肤",
        ],
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    print("=== Map Style L1 Feasibility ===")
    print("level:", verdict["level"])
    print("summary:", verdict["summary"])
    print(
        f"offline: {verdict['offline_pass']}/{verdict['offline_total']} | "
        f"retrieval_target: {verdict['retrieval_pass_target']}/{verdict['retrieval_total']} | "
        f"style_headings: {verdict['style_related_heading_count']}"
    )
    if retrieval_error:
        print("retrieval_error:", retrieval_error)
    for c in coverages:
        print(f"- {c.file}: paras={c.nonempty_paras} headings={c.headings_total} "
              f"style_h={len(c.style_related_headings)} codeish={c.codeish_paras}")
    for case in merged:
        off = case.get("offline") or {}
        ret = case.get("retrieval") or {}
        print(
            f"  [{case['id']}] offline={'PASS' if off.get('pass') else 'FAIL'} "
            f"retr_target={'PASS' if ret.get('pass_target_manual') else ('FAIL' if ret else 'SKIP')} "
            f"| {case['question'][:40]}"
        )
    print("report:", args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
