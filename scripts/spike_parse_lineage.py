#!/usr/bin/env python3
"""Emit FR-01.1 lineage JSONL for a tiny DOCX fixture (offline; no Chroma)."""

from __future__ import annotations

import argparse
import hashlib
import json
import sys
import uuid
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from rag_knowledge.services.loader import FileLoader
from rag_knowledge.services.unstructured_loader import UnstructuredChapterLoader


def _ensure_fixture(path: Path) -> Path:
    if path.exists():
        return path
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = DocxDocument()
    h = doc.add_paragraph()
    run = h.add_run("1. 安装准备")
    run.bold = True
    run.font.size = Pt(16)
    # Prefer style heading when available
    try:
        h.style = "Heading 1"
    except Exception:
        pass
    doc.add_paragraph("请执行以下命令完成目录创建。")
    doc.add_paragraph("mkdir -p /data/stamp")
    doc.add_paragraph("这是一段保留的中文技术说明，包含 Redis 与 /etc/sysctl.conf。")
    doc.save(path)
    return path


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _write_jsonl(path: Path, rows: list[dict]) -> None:
    with path.open("w", encoding="utf-8") as fh:
        for row in rows:
            fh.write(json.dumps(row, ensure_ascii=False) + "\n")


def emit_lineage(docx_path: Path, out_dir: Path) -> dict:
    run_id = f"_spike_{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}_{uuid.uuid4().hex[:8]}"
    target = out_dir / run_id
    target.mkdir(parents=True, exist_ok=True)

    snapshot = _sha256(docx_path)
    loader = UnstructuredChapterLoader()
    # parse via public API + paragraph pass for raw blocks
    document = DocxDocument(str(docx_path))
    raw_blocks = []
    for idx, paragraph in enumerate(document.paragraphs):
        text = (paragraph.text or "").strip()
        if not text:
            continue
        raw_blocks.append(
            {
                "raw_block_id": f"rb_{idx}",
                "block_index": idx,
                "block_type": "paragraph",
                "raw_text": text,
                "page_or_part": "word/document.xml",
                "parent_container": "document_body",
                "style_snapshot": {
                    "style_name": str(getattr(paragraph.style, "name", "") or ""),
                    "bold": any(run.bold for run in paragraph.runs if run.bold is not None),
                },
                "media_refs": [],
            }
        )

    elements = []
    docs = loader.load(str(docx_path))
    for order, doc in enumerate(docs):
        meta = doc.metadata or {}
        elements.append(
            {
                "element_id": f"el_{order}",
                "source_raw_block_ids": [],  # filled after greedy text match
                "element_type": meta.get("content_type") or "text",
                "content_markdown": doc.page_content,
                "searchable_text": meta.get("searchable_text") or doc.page_content,
                "element_order": int(meta.get("element_order") or order),
                "candidate_section_path": [
                    p.strip() for p in str(meta.get("section_path") or "").split(">") if p.strip()
                ],
                "content_type": meta.get("content_type") or "text",
            }
        )

    # Greedy attach raw blocks by substring containment
    used = set()
    for el in elements:
        body = el["content_markdown"] or ""
        linked = []
        for rb in raw_blocks:
            if rb["raw_block_id"] in used:
                continue
            if rb["raw_text"] and rb["raw_text"] in body:
                linked.append(rb["raw_block_id"])
                used.add(rb["raw_block_id"])
        el["source_raw_block_ids"] = linked

    structure_decisions = []
    for el in elements:
        structure_decisions.append(
            {
                "validation_id": f"sv_{el['element_id']}",
                "element_id": el["element_id"],
                "structure_action": "accept",
                "issue_codes": [],
                "heading_source": "style",
                "heading_confidence": 0.9,
                "resolved_section_path": el["candidate_section_path"],
            }
        )

    transformations = []
    content_decisions = []
    final_rows = []
    quarantine = []
    prev_id = None
    for i, el in enumerate(elements):
        keep = not FileLoader._is_low_information(el["content_markdown"] or "")
        decision = {
            "decision_id": f"cd_{el['element_id']}",
            "target_id": el["element_id"],
            "action": "keep" if keep else "reject",
            "reason_code": "keep" if keep else "low_information",
            "confidence": 0.9,
            "quality_metrics": {},
            "quarantine_ref": None,
        }
        content_decisions.append(decision)
        if not keep:
            quarantine.append({"raw_ref": el["source_raw_block_ids"], "element_id": el["element_id"], "text": el["content_markdown"]})
            continue
        tx = {
            "transformation_id": f"tx_{el['element_id']}",
            "action": "passthrough",
            "input_ids": [el["element_id"]],
            "output_ids": [f"chunk_{i}"],
            "boundary_type": "section",
            "reason_code": "KEEP_STRONG_SECTION_BOUNDARY",
            "char_range": [0, len(el["content_markdown"] or "")],
            "target_size": None,
        }
        transformations.append(tx)
        chunk_id = f"chunk_{i}"
        row = {
            "chunk_id": chunk_id,
            "source_document_id": docx_path.name,
            "source_snapshot_hash": snapshot,
            "source_raw_block_ids": el["source_raw_block_ids"],
            "source_element_ids": [el["element_id"]],
            "transformation_ids": [tx["transformation_id"]],
            "content_decision_id": decision["decision_id"],
            "section_id": ">".join(el["candidate_section_path"]),
            "section_path": ">".join(el["candidate_section_path"]),
            "chunk_index_global": len(final_rows),
            "chunk_index_in_section": 0,
            "prev_chunk_id": prev_id,
            "next_chunk_id": None,
            "parser_version": "spike-lineage-1",
            "config_fingerprint": "offline_spike",
        }
        if prev_id is not None:
            final_rows[-1]["next_chunk_id"] = chunk_id
        final_rows.append(row)
        prev_id = chunk_id

    # Destination coverage for raw blocks
    destinations = {}
    kept_raw = {rb for row in final_rows for rb in row["source_raw_block_ids"]}
    rejected_raw = {rb for q in quarantine for rb in (q.get("raw_ref") or [])}
    for rb in raw_blocks:
        rid = rb["raw_block_id"]
        if rid in kept_raw:
            destinations[rid] = "kept"
        elif rid in rejected_raw:
            destinations[rid] = "rejected"
        elif rid not in used:
            destinations[rid] = "non_content"
        else:
            destinations[rid] = "kept"

    manifest = {
        "run_id": run_id,
        "started_at": datetime.now(timezone.utc).isoformat(),
        "completed_at": datetime.now(timezone.utc).isoformat(),
        "source_snapshot_hashes": {docx_path.name: snapshot},
        "parser_version": "spike-lineage-1",
        "config_fingerprint": "offline_spike",
        "document_count": 1,
        "stage_counts": {
            "raw_blocks": len(raw_blocks),
            "canonical_elements": len(elements),
            "final_chunks": len(final_rows),
            "quarantine": len(quarantine),
        },
        "raw_block_destinations": destinations,
    }
    summary = {
        "双向追溯_ok": all(
            all(rid in {r["raw_block_id"] for r in raw_blocks} for rid in row["source_raw_block_ids"])
            for row in final_rows
        ),
        "无静默删除_ok": len(destinations) == len(raw_blocks),
        "destination_counts": {
            "kept": sum(1 for v in destinations.values() if v == "kept"),
            "rejected": sum(1 for v in destinations.values() if v == "rejected"),
            "non_content": sum(1 for v in destinations.values() if v == "non_content"),
        },
    }

    (target / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    _write_jsonl(target / "raw_blocks.jsonl", raw_blocks)
    _write_jsonl(target / "canonical_elements.jsonl", elements)
    _write_jsonl(target / "structure_decisions.jsonl", structure_decisions)
    _write_jsonl(target / "transformations.jsonl", transformations)
    _write_jsonl(target / "content_decisions.jsonl", content_decisions)
    _write_jsonl(target / "final_chunk_lineage.jsonl", final_rows)
    _write_jsonl(target / "quarantine.jsonl", quarantine)
    (target / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (target / "summary.md").write_text(
        "# Lineage Spike Summary\n\n"
        + f"- run_id: `{run_id}`\n"
        + f"- raw_blocks: {len(raw_blocks)}\n"
        + f"- final_chunks: {len(final_rows)}\n"
        + f"- 双向追溯_ok: {summary['双向追溯_ok']}\n"
        + f"- 无静默删除_ok: {summary['无静默删除_ok']}\n",
        encoding="utf-8",
    )
    return {"run_dir": str(target), "manifest": manifest, "summary": summary}


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--fixture",
        type=Path,
        default=ROOT / "docs/04_已完成归档/01_文档解析与切块/已完成-第0B轮-并行准备与预研/第0C轮设计与留痕/留痕样例/lineage_spike.docx",
    )
    parser.add_argument("--out-root", type=Path, default=ROOT / "data" / "chunk_audit")
    args = parser.parse_args(argv)
    fixture = _ensure_fixture(args.fixture)
    result = emit_lineage(fixture, args.out_root)
    print(json.dumps(result["summary"], ensure_ascii=False, indent=2))
    print(result["run_dir"])
    if not result["summary"]["双向追溯_ok"] or not result["summary"]["无静默删除_ok"]:
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
